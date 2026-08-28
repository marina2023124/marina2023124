#!/usr/bin/env python3
"""Generate bilingual income certificate Word template for visa applications."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path


def set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph(doc, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_mixed_paragraph(doc, parts, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    """parts: list of (text, bold, size)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    for text, bold, size in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_placeholder_field(doc, label_cn, label_en, width_cm=14):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F5F5F5")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label_cn} / {label_en}：")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run2 = p.add_run("【请填写 / Please fill in】")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run2.font.name = "Times New Roman"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "CCCCCC")
        borders.append(el)
    tblPr.append(borders)
    table.columns[0].width = Cm(width_cm)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build_document():
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Logo placeholder
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_p.add_run("【公司 LOGO / Company LOGO】")
    logo_run.font.size = Pt(9)
    logo_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    logo_run.font.name = "Times New Roman"
    logo_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_paragraph(
        doc,
        "北京寸阴成器科技有限公司",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "Beijing Cunyinchengqi Technology Co., Ltd.",
        bold=True,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_paragraph(
        doc,
        "【公司地址 / Company Address：请填写】",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    add_paragraph(
        doc,
        "Tel: 【电话】    Email: 【邮箱】",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    add_paragraph(doc, "收入证明", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(
        doc,
        "INCOME CERTIFICATE",
        bold=True,
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=16,
    )

    add_mixed_paragraph(
        doc,
        [
            ("兹证明 ", False, 11),
            ("【姓名】", True, 11),
            ("（证件号：", False, 11),
            ("【证件号】", True, 11),
            ("）为我公司员工，现任 ", False, 11),
            ("【部门】", True, 11),
            (" 部门 ", False, 11),
            ("【岗位】", True, 11),
            (" 岗位，自 ", False, 11),
            ("【入职日期，如 2020年3月15日】", True, 11),
            (" 起在我公司工作至今。", False, 11),
        ],
        space_after=8,
    )

    add_mixed_paragraph(
        doc,
        [
            (
                "This is to certify that ",
                False,
                11,
            ),
            ("【Full Name】", True, 11),
            (" (ID/Passport No.: ", False, 11),
            ("【ID/Passport No.】", True, 11),
            ("), is an employee of our company, currently serving as ", False, 11),
            ("【Position】", True, 11),
            (" in the ", False, 11),
            ("【Department】", True, 11),
            (" Department, and has been employed by our company since ", False, 11),
            ("【Start Date, e.g. March 15, 2020】", True, 11),
            (".", False, 11),
        ],
        space_after=12,
    )

    add_mixed_paragraph(
        doc,
        [
            ("该员工目前税前月薪为人民币 ", False, 11),
            ("【税前月薪，如 25,000】", True, 11),
            (" 元（大写：", False, 11),
            ("【月薪大写，如 贰万伍仟元整】", True, 11),
            ("），税前年薪约为人民币 ", False, 11),
            ("【税前年薪，如 300,000】", True, 11),
            (" 元。", False, 11),
        ],
        space_after=8,
    )

    add_mixed_paragraph(
        doc,
        [
            (
                "The employee's current pre-tax monthly salary is RMB ",
                False,
                11,
            ),
            ("【Monthly Salary, e.g. 25,000】", True, 11),
            (" (in words: ", False, 11),
            ("【Monthly Salary in Words】", True, 11),
            ("), with an approximate annual pre-tax income of RMB ", False, 11),
            ("【Annual Salary, e.g. 300,000】", True, 11),
            (".", False, 11),
        ],
        space_after=12,
    )

    add_paragraph(
        doc,
        "上述收入真实有效，特此证明。本证明仅用于办理签证（申根签证 / 美国签证等）之目的。",
        size=11,
        space_after=6,
    )
    add_paragraph(
        doc,
        "The above income information is true and accurate. This certificate is issued solely for visa application purposes (including Schengen visa, U.S. visa, etc.).",
        size=11,
        space_after=12,
    )

    add_paragraph(doc, "— 员工信息摘要 / Employee Information Summary —", bold=True, size=10, space_after=8)

    fields = [
        ("姓名", "Name"),
        ("证件号", "ID / Passport No."),
        ("部门", "Department"),
        ("岗位", "Position"),
        ("入职日期", "Date of Employment"),
        ("税前月薪（元）", "Pre-tax Monthly Salary (RMB)"),
        ("税前年薪（元）", "Pre-tax Annual Salary (RMB)"),
    ]
    for cn, en in fields:
        add_placeholder_field(doc, cn, en)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in sig_table.rows:
        for cell in row.cells:
            for edge in ("top", "left", "bottom", "right"):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                borders = OxmlElement("w:tcBorders")
                for b in ("top", "left", "bottom", "right"):
                    el = OxmlElement(f"w:{b}")
                    el.set(qn("w:val"), "nil")
                    borders.append(el)
                tcPr.append(borders)

    left_cells = [
        ("授权签字人 / Authorized Signatory：", "【HR负责人姓名 / Name of HR Manager】"),
        ("职务 / Position：", "【职务，如人力资源经理 / HR Manager】"),
        ("签发日期 / Date of Issue：", "【签发日期，如 2026年8月28日 / August 28, 2026】"),
        ("公司盖章 / Company Seal：", "（此处加盖公章）\n(Official company seal here)"),
    ]
    for i, (label, value) in enumerate(left_cells):
        cell = sig_table.rows[i].cells[0]
        p = cell.paragraphs[0]
        r1 = p.add_run(label + "\n")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "Times New Roman"
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r2 = p.add_run(value)
        r2.font.size = Pt(10)
        r2.font.name = "Times New Roman"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for i in range(4):
        sig_table.rows[i].cells[1].text = ""

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note_run = note.add_run(
        "填写说明 / Instructions：\n"
        "1. 请将文中所有【】内占位内容替换为真实信息，删除灰色提示文字。\n"
        "   Replace all bracketed placeholders with actual information and remove gray hints.\n"
        "2. 申根签、美签通常要求英文内容完整；请确保英文姓名、职位、日期格式正确。\n"
        "   For Schengen and U.S. visas, ensure English name, position, and date formats are correct.\n"
        "3. 月薪大写示例：壹万贰仟元整 / Twenty Thousand Yuan Only。\n"
        "4. 打印后需加盖公司公章并由 HR 签字，再扫描或拍照上传至收入证明申请系统。\n"
        "   After printing, obtain company seal and HR signature before uploading to the HR system.\n"
        "5. 如 HR 系统要求上传附件，可将本文件作为自定义模板附件提交。\n"
        "   Upload this document as a custom attachment in the income certificate application."
    )
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    note_run.font.name = "Times New Roman"
    note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    return doc


def main():
    out_dir = Path("/workspace/documents")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "收入证明_Income_Certificate_签证用模板.docx"
    doc = build_document()
    doc.save(out_path)
    print(f"Generated: {out_path}")


if __name__ == "__main__":
    main()
